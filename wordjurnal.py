from docx import Document
from docx.shared import Inches, Mm

months = ['JANUARI', 'FEBUARI', 'MARET', 'APRIL', 'MEI', 'JUNI', 'JULI', 'AGUSTUS', 'SEPTEMBER', 'OKTOBER', 'NOVEMBER', 'DESEMBER']

document = Document()

p = document.add_paragraph()
r = p.add_run()
section = document.sections[0]
section.page_height = Mm(210)
section.page_width = Mm(148)
for y in range(1, 13):
    for x in range(1, 32, 3):
        try:
            r.add_picture(f'foto/{y}/{months[y-1]}{x}.png')
        except:
            pass

document.save('demo.docx')